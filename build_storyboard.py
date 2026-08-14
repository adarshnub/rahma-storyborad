from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'RAHMA_10s_Generation_Storyboard_Part_1.docx'
MD = ROOT / 'RAHMA_10s_Generation_Storyboard_Part_1.md'
REF = ROOT / 'reference-images'

NAVY = '0B1B33'; INDIGO = '14213D'; TEAL = '2E6F73'; GOLD = 'C9922C'; MIST = 'EEF1F4'; INK = '1D2630'; MUTED = '59636E'

base_prompt = """Use case: photorealistic-natural. Asset: 16:9, 10-second cinematic video clip for the Rahma launch film. Grade: muted indigo shadows, warm amber practicals, cream limestone, restrained teal; 24 fps, 180-degree shutter, gentle 35mm grain, natural skin texture. Camera language: intimate, human, precise, never a surveillance view. The camera stays with a person or a meaningful object. Motion is slow, stable and deliberate. Every character is Arab, from a contemporary UAE/GCC context. Performance: quiet attention, never fear or spectacle. No sirens, no flashing lights, no ambulance action, no panic, no crowd, no unapproved visible brands, no readable incidental text, no watermark. The Rahma phone is graphite black with a deep-indigo screen and one large round amber-gold hold button marked bold uppercase RAHMA in deep indigo; the round black wearable shows bold uppercase RAHMA in warm off-white inside a thin amber ring. Preserve the supplied reference frames exactly for wardrobe, locations, device finish and palette."""

clips = [
    dict(n='01', tc='00:00-00:10', edit='Use 00:00-00:10.', refs='None (licensed archive only; use the archive fallback note below).',
         beat='0:00-0:03 absolute black. 0:03-0:10 grainy 1971 signing footage: tight hands, fountain pen, formal document, slight slow motion. End with the pen completing the signature and lifting.',
         vo='This country did not begin with a border.', sound='0:00-0:03 true silence. At 0:03 a single low sustained cello or oud note enters; no other effects.',
         prompt='Begin on full black for exactly three seconds. Then reveal a historically respectful, grainy 1971 archival-style macro: mature male hands in a dark formal sleeve signing an official founding document with a fountain pen. 50mm close shot, soft monochrome-sepia archive texture, subtle film gate weave, slowed just enough to feel remembered. Do not show a wide ceremony, flags, crowds, logos, or legible document text. End on the pen lifting after the completed signature.',
         first='Pure black.', last='Tight archival frame: pen just lifted from the completed signature.', handoff='Hard cut on the pen lift to archival faces in Clip 02; retain the same monochrome-sepia grade for the first beat.'),
    dict(n='02', tc='00:10-00:20', edit='Use 00:10-00:20.', refs='REF-02 for tea-hands texture only.',
         beat='0:10-0:14 archival faces of founding fathers, intimate rather than ceremonial. 0:14-0:18 slow push toward Sheikh Zayed archive portrait. 0:18-0:20 present-day: school gate, hospital corridor, older hands around tea; three brisk 0.7-second human details.',
         vo='It began with a promise. That no one here would be left to manage alone. Everything built since has been that promise, keeping its word.', sound='The sustained note opens to restrained warm strings at 0:18. Keep the archive quiet; no crowd ambience.',
         prompt='Open on a monochrome-sepia archival close composition of founding-era faces, not a ceremonial wide. Make a gentle 4-second push toward a respectful archival portrait of Sheikh Zayed. Then make three clean, human present-day cuts: children passing through a modest UAE school gate at 7am; a quiet clean hospital corridor; lined older hands around a clear glass cup of tea. Each contemporary shot is intimate, 35mm, warm and unhurried. No skyline, flags, legible signage, visible brands or theatrical emotion.',
         first='Archive face frame, close and grainy.', last='Lined hands still around a tea glass in warm morning light.', handoff='Hard cut to black/type in Clip 03. The tea warmth becomes the gold accent of the typography.'),
    dict(n='03', tc='00:20-00:30', edit='Use 00:20-00:30.', refs='No photoreal reference. Use approved bilingual graphic artwork for exact Arabic/English text.',
         beat='0:20-0:22 black hold. 0:22-0:29 first typographic card, Arabic above English: قانون حقوق الطفل / A law for its children. 0:29-0:30 black dip, anticipating the next card.',
         vo='For its children, it wrote a law - and made it cover every child in this country. Not only its own.', sound='Warm strings continue, understated. No whoosh on the text.',
         prompt='Create a controlled graphic video, not a 3D logo animation. Matte black background. At 2 seconds, fade in approved, compositor-supplied Arabic lettering above a smaller English line, both centered, warm off-white, elegant and still. Hold without particle effects, flares, movement or UI motifs; dip gently to black in the final second. IMPORTANT: do not ask the generator to render text. Composite the exact supplied vector text in post: Arabic "قانون حقوق الطفل" and English "A law for its children".',
         first='Black, held.', last='Black after the first title has fully cleared.', handoff='Clip 04 begins from identical black. Keep font, leading, scale and warm off-white color pixel-identical.'),
    dict(n='04', tc='00:30-00:40', edit='Use 00:30-00:40.', refs='No photoreal reference. Use approved bilingual graphic artwork for exact Arabic/English text.',
         beat='0:30-0:34 card 2: أصحاب الهمم / People of Determination. 0:34-0:38 card 3: كبار المواطنين / Senior Citizens. 0:38-0:40 clear to black and hold.',
         vo='For people with disabilities, it changed the word. People of Determination. For the elderly, it changed the word again. Senior Citizens. A nation is measured by what it calls the people it protects.', sound='Strings remain warm and restrained; fade down toward one held note by 0:40.',
         prompt='On matte black, present two sequential calm title cards in the exact same approved type system as Clip 03. First card holds for four seconds; fade to black for a brief breath; second card holds for four seconds; end on two seconds of full black. No animated streaks, lens flares, sound-design whooshes, logos, icons or texture. Composite exact vectors in post; do not rely on video AI for words: card 1 Arabic "أصحاب الهمم", English "People of Determination"; card 2 Arabic "كبار المواطنين", English "Senior Citizens".',
         first='Identical matte black to Clip 03 end.', last='Full black held two seconds.', handoff='Carry black forward. Clip 05 introduces the Rahma word from this exact darkness.'),
    dict(n='05', tc='00:40-00:50', edit='Use 00:40-00:50.', refs='Approved Rahma Arabic wordmark vector; do not generate the lettering.',
         beat='0:40-0:42 full black, music almost gone. 0:42-0:45 Arabic رحمة appears alone, large and centred. 0:45-0:47 hold in silence. 0:47-0:50 small English Rahma fades beneath.',
         vo='There is one more word. Rahma. It is not a word this country borrowed. It is the one it was built on.', sound='One held note falls away at 0:42. Maintain two full seconds of silence under the Arabic word.',
         prompt='Minimal graphic film clip on a pure matte black background. Keep black for 2 seconds. Fade in only the approved Arabic Rahma wordmark, centered and large in warm off-white; hold absolutely still for 2 seconds in silence. Then fade in a small English "Rahma" beneath it. No logo reveal effects, no flare, no texture, no glow beyond clean anti-aliased type. Composite supplied vector artwork in post; do not use generated text.',
         first='Full black.', last='Arabic wordmark and small English name, centered and motionless on black.', handoff="Match-cut the Arabic wordmark's final silhouette into the app mark at the start of Clip 06; preserve the warm off-white-to-amber transition."),
    dict(n='06', tc='00:50-01:00', edit='Use 00:50-01:00.', refs='REF-02 Home Environments; REF-04 Device and Wearable.',
         beat='0:50-0:53 the Arabic word transforms slowly into the Rahma app mark (composite). 0:53-1:00 cut to a graphite phone waking in the pre-dawn bedroom, held by Mariam; the screen settles on one large amber hold button marked bold RAHMA.',
         vo='And now it is something you can hold in your hand.', sound='At 0:53 warm strings re-enter softly. No digital whoosh; use only a soft, tactile wake chime.',
         prompt='Start on the clean app-mark silhouette from the end of Clip 05; this 3-second vector morph will be composited, slow and quiet. Match cut to Mariam in the pre-dawn bedroom from REF-02. In a close over-shoulder shot, her lined hands lift the graphite black phone from the walnut bedside table. The deep-indigo display wakes calmly and resolves to a single large round amber-gold hold button with a faint teal halo, marked bold uppercase RAHMA in deep indigo. Camera drifts forward a few centimetres, 50mm, one warm lamp and blue dawn through sheer curtains. Use REF-04 exactly for device shape and UI. No unreadable UI clutter, no notification barrage, no other people.',
         first='Centred app-mark silhouette on black.', last='Mariam holds the awake phone close to camera; amber button centred, her thumbs resting but not pressing.', handoff='Clip 07 begins on the same phone, same hands, same upright orientation. Match the button position and lamp/dawn colour.'),
    dict(n='07', tc='01:00-01:10', edit='Use 01:00-01:10.', refs='REF-01 Cast Continuity (Mariam); REF-02 Home Environments; REF-04 Device and Wearable.',
         beat="Mariam's lined, ringed hands fill the frame. Her thumb presses and holds the single button. Count one, two. At the end of second two, a soft teal confirmation pulse resolves. Continue holding on the calm result; never show menu navigation.",
         vo='Rahma was built for the people who find technology hardest. No typing. No menus. One button, held for two seconds.', sound='Close room tone. A nearly inaudible tactile confirmation at second two; music remains gently lifted.',
         prompt='Continue directly from Clip 06: same Mariam, same dove-grey headscarf and cream cardigan, same lined hands and gold wedding band, same graphite phone and deep-indigo UI. Macro 85mm in the calm pre-dawn bedroom. Her thumb lowers on the single amber circle marked bold uppercase RAHMA and deliberately holds for exactly two seconds. Count the hold through stillness, not graphics. At the end of second two, the circle gives one soft teal confirmation pulse; her hand relaxes but stays in frame. She is capable and calm, not confused. No typing, menus, app text, medical alarms, urgency or extra fingers.',
         first="Exact continuation: phone upright in Mariam's two hands, button unpressed.", last='The soft teal confirmation pulse has settled into a quiet amber circle; her steady hand remains in frame.', handoff='Use the circle pulse as a graphic match-cut to the location marker glow in Clip 08. Do not change the indigo/amber/teal palette.'),
    dict(n='08', tc='01:10-01:20', edit='Use 01:10-01:20.', refs='REF-01 Cast Continuity (Lina and Sami); REF-03 Service Locations; REF-04 Device and Wearable.',
         beat="1:10-1:15 over Lina's shoulder in her parked graphite car at dusk: a simplified, non-readable map resolves to building and floor. 1:15-1:20 Sami is mid-conversation indoors; phone lights, he stops, and rises with quiet certainty.",
         vo='Rahma knows where that citizen is - even when the signal does not. And it tells the people who love them, first.', sound="A soft map-resolution tone bridges the match cut. Sami's room tone stops naturally as he stands; do not use an alert siren.",
         prompt='Begin with the teal pulse from Clip 07 becoming a simple amber location dot on a non-readable deep-indigo map seen over the shoulder of Lina, an Arab caregiver from the UAE/GCC seen in REF-01, seated calmly in the graphite compact car from REF-03 on a quiet UAE residential street at dusk. Show only abstract building and floor shapes, no readable addresses. After five seconds, cut to Sami, an Arab man from REF-01, in a warm apartment conversation; his graphite phone lights softly on a nearby table, he looks, pauses, and stands before anyone else reacts. Camera stays near him, not on the screen. Understated, warm, confident, no driving, no running, no panic.',
         first='Teal pulse expands into an amber map marker over a deep-indigo map.', last='Sami just fully upright, phone in hand, face composed and attentive.', handoff="Clip 09 opens on Sami's attentive look and expands the notification chain. Preserve his brown overshirt, cream T-shirt and phone finish."),
    dict(n='09', tc='01:20-01:30', edit='Use 01:20-01:30.', refs='REF-01 Cast Continuity (Sami, Dr. Noor, control operator); REF-03 Service Locations; REF-04 Device and Wearable.',
         beat='1:20-1:21.5 Sami begins to move out of frame. 1:21.5-1:23 doctor sees a quiet screen and turns to walk. 1:23-1:24.5 operator lifts her head in the control room. 1:24.5-1:26 paramedic closes an already-open rear vehicle door, no lights/siren. 1:26-1:30 a woman reads a phone: relief, not panic.',
         vo='Then it reaches everyone else who needs to know. Doctors. Police. Ambulance. The people who care for them at home. All at once. Without anyone having to explain. It carries what a doctor would need to know before they arrive.', sound='Music gains gentle momentum but never percussion-heavy. Each acknowledgement has a soft, related tactile tone. Absolutely no siren.',
         prompt='Open on Sami from Clip 08 already moving with composed purpose. Make four humane 1.5-second match-cut vignettes: Dr. Noor from REF-01 glances at a non-readable phone and turns into the quiet hospital corridor from REF-03; the navy-uniform control operator from REF-01 lifts her head at her desk, screens remain abstract and dark; a non-speaking paramedic calmly closes a rear vehicle door with no flashing light; then a mid-40s Arab woman in a warm apartment reads her graphite phone, her face releases into relief. Every screen remains secondary and unreadable; all humans are attentive, never frantic. No vehicle movement, no ambulance branding, no emergency spectacle.',
         first='Sami in brown overshirt, already turning out of his conversation.', last="Relieved woman's face, phone lowered slightly, breathing easy.", handoff='Cut from her relaxed face to the quiet wearable macro in Clip 10; keep emotional temperature low and safe.'),
    dict(n='10', tc='01:30-01:40', edit='Generate 10 seconds; use only 01:30-01:36 and fade to black. Do not use the final four seconds.', refs='REF-01 Cast Continuity (Mariam); REF-02 Home Environments; REF-04 Device and Wearable.',
         beat='1:30-1:33 thin wrist on white bedsheet, wearable resting in morning light. 1:33-1:36 pull back to Mariam quietly reading the newspaper in her cream-and-walnut home; nothing is happening. 1:36 cut to black. The extra generated 1:36-1:40 is discarded.',
         vo='It speaks Arabic. English. And more. It remembers the medicine, when the day gets long. And when the network fails, it keeps going anyway. Quietly. In the background. Before anything goes wrong.', sound='Music thins to a soft single line by 1:36, with natural paper rustle. No added cues after the edit-out.',
         prompt="Begin on the exact round black wearable from REF-04 resting on Mariam's thin wrist over white cotton bedding in soft morning light. Its deep-indigo face carries a thin amber ring with bold uppercase RAHMA centred in warm off-white; do not animate a data dashboard. Over six calm seconds, pull back from the wrist to reveal Mariam from REF-01 in the cream-and-walnut home from REF-02, seated comfortably and reading a folded newspaper. She is in her dove-grey headscarf, muted teal blouse and cream cardigan. Nothing has happened; that is the point. Warm domestic quiet, 50mm, subtle grain. Hold four extra seconds of the same unremarkable peace for editorial safety, then discard them. No alerts, illness, drama, nurses, logos or busy UI.",
         first='Wearable macro: amber ring quiet on white bedding.', last='EDITORIAL OUT at 01:36: Mariam reading in peaceful morning light. Cut to black before the next section begins.', handoff='Do not generate or storyboard "Three Moments" here. Editorial hard cut to black at 01:36.'),
]

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths): cell.width = Inches(width)
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'),'dxa')
    tblInd = tblPr.first_child_found_in('w:tblInd')
    if tblInd is None: tblInd = OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths): col.set(qn('w:w'), str(round(width*1440)))

def font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = 'Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); run.bold=bold; run.italic=italic

def p(doc, text='', style=None, size=None, color=None, bold=False, italic=False, after=None, before=None, keep=False):
    para=doc.add_paragraph(style=style)
    if text:
        r=para.add_run(text); font(r, size or (11 if style is None else 11), color or INK, bold, italic)
    pf=para.paragraph_format
    if after is not None: pf.space_after=Pt(after)
    if before is not None: pf.space_before=Pt(before)
    if keep: pf.keep_with_next=True
    return para

def label_para(doc, label, text, after=3):
    para=doc.add_paragraph(); para.paragraph_format.space_after=Pt(after); para.paragraph_format.line_spacing=1.12
    r=para.add_run(label+'  '); font(r, 9.5, GOLD, True)
    r=para.add_run(text); font(r, 10, INK)
    return para

def add_page_break(doc): doc.add_page_break()

def add_header_footer(section):
    header=section.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=header.add_run('RAHMA | PART 1 GENERATION STORYBOARD'); font(r,8.5,MUTED,True)
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=footer.add_run('Confidential production working document'); font(r,8,MUTED)

def add_ref_page(doc, filename, title, caption):
    p(doc, title, 'Heading 1', after=6, keep=True)
    p(doc, caption, size=10.5, color=MUTED, after=12)
    doc.add_picture(str(REF/filename), width=Inches(6.25))
    cap=p(doc, 'Generated visual reference - use as continuity guidance, not final footage.', size=8.5, color=MUTED, italic=True, after=0)
    add_page_break(doc)

def add_clip(doc, c):
    p(doc, f"CLIP {c['n']} | {c['tc']}", 'Heading 1', after=2, keep=True)
    p(doc, c['edit'], size=10, color=GOLD, bold=True, after=8)
    label_para(doc,'EDITORIAL BEAT',c['beat'],after=5)
    label_para(doc,'VOICEOVER',c['vo'],after=5)
    label_para(doc,'SOUND / MUSIC',c['sound'],after=5)
    label_para(doc,'REFERENCE INPUTS',c['refs'],after=8)
    p(doc,'VIDEO-GENERATION PROMPT', 'Heading 2', after=4, keep=True)
    prompt_box=doc.add_table(rows=1, cols=1); prompt_box.alignment=WD_TABLE_ALIGNMENT.LEFT; set_table_widths(prompt_box,[6.5])
    cell=prompt_box.cell(0,0); set_cell_shading(cell,MIST); set_cell_margins(cell,120,160,120,160); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    pp=cell.paragraphs[0]; pp.paragraph_format.space_after=Pt(0); pp.paragraph_format.line_spacing=1.08
    r=pp.add_run(base_prompt+' '+c['prompt']); font(r,9.2,INK)
    p(doc,'CONTINUITY HANDOFF', 'Heading 2', after=4, before=10, keep=True)
    label_para(doc,'FIRST FRAME',c['first'],after=3)
    label_para(doc,'LAST FRAME',c['last'],after=3)
    label_para(doc,'NEXT-CLIP BRIDGE',c['handoff'],after=0)
    add_page_break(doc)

def build_doc():
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(0.492); sec.footer_distance=Inches(0.492)
    add_header_footer(sec)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name, sz, col, before, after in [('Heading 1',16,NAVY,18,10),('Heading 2',13,TEAL,14,7),('Heading 3',12,INDIGO,10,5)]:
        s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size=Pt(sz); s.font.color.rgb=RGBColor.from_string(col); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    # Cover: named override to compact_reference_guide uses Rahma indigo/amber instead of preset blue.
    p(doc,'RAHMA', size=12, color=GOLD, bold=True, after=10)
    title=p(doc,'10-Second Generation Storyboard', size=28, color=NAVY, bold=True, after=5); title.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p(doc,'Part 1: The Promise, The Word, and What Rahma Can Do | 00:00-01:36', size=13, color=MUTED, after=22)
    lead=p(doc,'A continuity-first production pack for generating connected 10-second video clips. It deliberately stops before the “Three Moments” section.', size=12, color=INK, after=20)
    t=doc.add_table(rows=4, cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT; set_table_widths(t,[1.8,4.7])
    data=[('Source','Rahma Film - Shooting Script, Draft 1'),('Coverage','00:00-01:36 only; 10 generated clips'),('Generation standard','16:9 | 24 fps | 10 seconds per source clip'),('Edit standard','Clip 10 is generated to 10 seconds but cut at 01:36')]
    for row,(a,b) in zip(t.rows,data):
        for cell in row.cells: set_cell_margins(cell)
        set_cell_shading(row.cells[0],MIST); r=row.cells[0].paragraphs[0].add_run(a); font(r,9.5,NAVY,True); r=row.cells[1].paragraphs[0].add_run(b); font(r,9.5,INK)
    p(doc,'Working intent', 'Heading 1', after=5)
    p(doc,'The film makes attention visible without turning care into spectacle. Every human reaction is composed; every device is secondary to a person; every transition has an explicit first/last-frame handoff so separate generations can cut together cleanly.', size=11, after=0)
    add_page_break(doc)

    p(doc,'How to use this pack','Heading 1',after=6)
    for txt in [
        'Load the named reference image(s) with every applicable generation. Do not introduce new wardrobe, device finishes, UI layouts, locations, or casting.',
        'Generate the entire 10 seconds for every clip. Preserve a clean first and last 12 frames; editorial uses these frames for the stated transition.',
        'Build all Arabic/English cards, the Rahma word and app-mark morph in compositing with approved vector artwork. Do not ask a video model to create legible Arabic or English text.',
        'Treat "no sirens" as a picture and sound rule. No flashing red/blue lights, rushing, crowded rooms, panic performance, or rescue iconography.',
        'Use 24 fps, 16:9 and the same grade across all clips. Keep faces natural, domestic details tactile, and screen exposure at least one stop below facial exposure.'
    ]:
        para=doc.add_paragraph(style='List Bullet'); para.paragraph_format.space_after=Pt(4); r=para.add_run(txt); font(r,10.5,INK)
    p(doc,'Timing map','Heading 1',after=6)
    p(doc,'The master cut ends Part 1 at 01:36. Each source generation is ten seconds; Clip 10 has four seconds of editorial overrun that are intentionally discarded.',size=10.5,color=MUTED,after=8)
    timing=doc.add_table(rows=1,cols=4); timing.alignment=WD_TABLE_ALIGNMENT.LEFT; set_table_widths(timing,[0.75,1.15,1.55,3.05])
    for cell,txt in zip(timing.rows[0].cells,['Clip','Source','Master use','Purpose']):
        set_cell_shading(cell,NAVY); set_cell_margins(cell); r=cell.paragraphs[0].add_run(txt); font(r,9,'FFFFFF',True)
    purposes=['Archive signature','Promise / present-day proof','Child-law card','Respectful naming cards','Rahma word','Word to hand-held app','One-button hold','Location and loved-one notice','Care network notices','Wearable / preventive quiet']
    for c,purpose in zip(clips,purposes):
        row=timing.add_row(); vals=[c['n'],'10 sec',c['tc'] if c['n']!='10' else '01:30-01:36','%s'%purpose]
        for cell,txt in zip(row.cells,vals): set_cell_margins(cell); r=cell.paragraphs[0].add_run(txt); font(r,9,INK)
    add_page_break(doc)

    p(doc,'Continuity bible','Heading 1',after=6)
    p(doc,'These are locked anchors. If a desired shot conflicts with one, preserve the anchor and change the shot.',size=10.5,color=MUTED,after=10)
    p(doc,'Characters','Heading 2',after=4)
    for text in [
        'Mariam - 74, Arab woman; dove-grey headscarf, muted teal blouse, cream cardigan, cream trousers, fine gold wedding band. Calm, competent, never frail-coded.',
        'Lina - late-30s Arab caregiver from the UAE/GCC; dark hair covered in a light-stone hijab, sand-beige blouse, dark navy trousers. Uses the graphite compact car at dusk.',
        'Sami - 27, Arab man; short black hair, brown overshirt over cream T-shirt, dark trousers. Receives a notification mid-conversation and stands without drama.',
        'Dr. Noor - 40s Arab woman; dark-teal scrubs under white coat. Purposeful, not rushed.',
        'Control operator - 30s Arab woman in a navy hijab and matte navy uniform. Alert but composed. The prompt sheet permits a non-speaking paramedic only as a short insert.'
    ]:
        para=doc.add_paragraph(style='List Bullet'); para.paragraph_format.space_after=Pt(3); r=para.add_run(text); font(r,10,INK)
    p(doc,'Objects, UI and spaces','Heading 2',after=4)
    for text in [
        'Phone: slim graphite-black slab, dark graphite edge, deep-indigo screen. One round amber-gold hold button marked bold uppercase RAHMA in deep indigo, plus restrained teal confirmation pulse. Never show a keyboard, dense menu or legible data.',
        'Wearable: round black case and black strap, deep-indigo face, thin amber ring with bold uppercase RAHMA in warm off-white. It is quiet; no dashboard animation.',
        'Homes: cream/off-white walls, walnut cabinetry or furniture, white cotton bedding, restrained teal accents, glass tea cup, kettle, small green plant, folded newspaper. Modest and lived-in, not luxurious or stylized.',
        'Care spaces: quiet cream-and-light-wood hospital corridor, dark matte-navy control room with abstract unreadable maps, low-rise UAE residential street at blue hour. No skyline or dramatic public-space scale.',
        'Lighting: black/sepia archive -> warm morning/cream domestic light -> deep indigo device screens with amber and teal accents. No red emergency palette.'
    ]:
        para=doc.add_paragraph(style='List Bullet'); para.paragraph_format.space_after=Pt(3); r=para.add_run(text); font(r,10,INK)
    p(doc,'Global negative prompt','Heading 2',after=4)
    p(doc,'No non-Arab casting, sirens, flashing emergency lights, ambulance chase, crowded public spaces, crying, falls, medical distress, surveillance camera angles, sci-fi UI, unapproved logos, unreadable AI text, skyline hero shots, flags, glamour styling, or stock-photo smiles.', size=10.2, color='7A2E2E', after=0)
    add_page_break(doc)

    add_ref_page(doc,'REF-01-cast-continuity.png','Reference atlas 01 | Cast continuity','Use these five people and their wardrobe anchors everywhere in Part 1. The line-up is an identity reference, not a scene composition.')
    add_ref_page(doc,'REF-02-home-environments.png','Reference atlas 02 | Home environments','Left: pre-dawn bedroom for Clip 06. Right: Mariam’s morning home for Clip 10. Keep the walnut, cream, teal and quiet domestic scale.')
    add_ref_page(doc,'REF-03-service-locations.png','Reference atlas 03 | Service locations','Left: Lina’s parked car at dusk. Centre: hospital corridor. Right: control room. These spaces must remain restrained and non-spectacular.')
    add_ref_page(doc,'REF-04-device-and-wearable.png','Reference atlas 04 | Device and wearable','Lock the graphite phone, deep-indigo UI, amber RAHMA button, teal pulse, round black wearable and amber RAHMA ring.')

    p(doc,'Clip-by-clip generation prompts','Heading 1',after=4)
    p(doc,'Every card contains the full generation constraint set, the exact editorial function, voiceover, sound, prompt, and handoff. Use each as a self-contained prompt while preserving the listed bridge.',size=10.5,color=MUTED,after=0)
    add_page_break(doc)
    for c in clips: add_clip(doc,c)
    p(doc,'Part 1 endpoint','Heading 1',after=6)
    p(doc,'At 01:36, cut Mariam’s ordinary morning to full black. This pack intentionally contains no material from “Three Moments” or the close.',size=11,after=6)
    p(doc,'Preflight before generation','Heading 2',after=4)
    for item in ['Approved/licensed 1971 archive or object-and-stills fallback for Clips 01-02.','Approved Arabic and English title artwork for Clips 03-05, plus Rahma app-mark vector for Clip 06.','Reference atlas images loaded into every applicable video-generation job.','One confirmed Arabic-first / English-subtitle VO approach before final timing lock.']:
        para=doc.add_paragraph(style='List Bullet'); para.paragraph_format.space_after=Pt(4); r=para.add_run(item); font(r,10.5,INK)
    doc.save(OUT)

def build_md():
    lines=['# RAHMA - 10-Second Generation Storyboard (Part 1)','', '**Scope:** 00:00-01:36 only. Stops before “Three Moments”. Every source clip is 10 seconds; Clip 10 is trimmed to 6 seconds in the master edit.','', '## Locked continuity','', '- Cast: every character is Arab, from a contemporary UAE/GCC context.','- Phone: graphite black; deep-indigo UI; one amber-gold round hold button marked bold uppercase RAHMA; subtle teal confirmation pulse.','- Wearable: black round case and strap; deep-indigo face; thin amber ring with bold uppercase RAHMA in warm off-white.','- Palette: muted indigo, cream limestone, walnut, restrained teal, amber practicals. No red emergency palette.','- Rule: no sirens, flashing lights, panic, crowds, rescue spectacle, unreadable AI text, or surveillance framing.','', '## Reference images','']
    for f,desc in [('REF-01-cast-continuity.png','Cast and wardrobe anchors'),('REF-02-home-environments.png','Pre-dawn bedroom and Mariam home'),('REF-03-service-locations.png','Car, hospital corridor, control room'),('REF-04-device-and-wearable.png','Phone UI and wearable')]: lines += [f'- `reference-images/{f}` - {desc}']
    lines += ['', '## Clips','']
    for c in clips:
        lines += [f"### Clip {c['n']} | {c['tc']}", '', f"**Editorial use:** {c['edit']}", '', f"**Beat:** {c['beat']}", '', f"**VO:** {c['vo']}", '', f"**Sound:** {c['sound']}", '', f"**References:** {c['refs']}", '', '**Prompt:**', '', base_prompt+' '+c['prompt'], '', f"**First frame:** {c['first']}", '', f"**Last frame:** {c['last']}", '', f"**Bridge:** {c['handoff']}", '']
    MD.write_text('\n'.join(lines), encoding='utf-8')

if __name__ == '__main__':
    build_doc(); build_md(); print(OUT); print(MD)
